from docx import Document
from docx.shared import RGBColor
import re

def hex_to_rgbcolor(hex_str):
    """#RRGGBB → RGBColor"""
    hex_str = hex_str.lstrip("#")
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)

def export_to_docx(all_messages, name_color, output_path):
    doc = Document()

    for msg in all_messages:
        name = msg["name"]
        content_lines = msg["content"]

        # 转换颜色
        color_hex = name_color.get(name, "#000000")  # 默认黑色
        color = hex_to_rgbcolor(color_hex)

        # 创建段落
        paragraph = doc.add_paragraph()

        # 添加 name
        run_name = paragraph.add_run(name)
        run_name.font.color.rgb = color

        # 添加 tab
        paragraph.add_run("\t")

        # 添加 content（多行合并为一段）
        text = " ".join(content_lines)
        run_text = paragraph.add_run(text)
        run_text.font.color.rgb = color

    doc.save(output_path)

